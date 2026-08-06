"""Build one .docx shell per resume template, into templates/resumes/.

    python templates/build_resume_docx_templates.py

Not part of the running app. Run it when a template's typography changes; the
resulting files are then treated as fixed shells that only ever get fields
filled into them, exactly as the single resume_template.docx was before.

WHY A REBUILD RATHER THAN AN EDIT
---------------------------------
The predecessor of this script produced one Georgia template, and that template
had a quiet bug: it emitted `{{ edu.school }}` and `{{ edu.year }}` while the
context that reaches it carries `institution`, `degree` and `year`. Every Word
export has therefore been shipping an Education section with a blank line where
the university should be. It was invisible because the PDF — which is what
almost everyone downloads — used a different, correct template. Rebuilding from
one description of the document fixes that by construction: both renderers now
read the same field names, in the same order, under the same headings.

WHAT DIFFERS PER TEMPLATE, AND WHAT CANNOT
------------------------------------------
Word has no reflowing tinted band and no cheap equivalent of several of the
finer print refinements, and forcing them in would mean tables or frames —
precisely the shapes that break an ATS. So a DOCX carries its template's
typography (family, sizes, weight, heading case, and a rule where the design
has one) and not its every flourish. Section order, headings and content are
identical to the PDF, which is what a parser reads in either format.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent / "resumes"

#: Mirrors services/resume_templates.TEMPLATES. Kept as plain data so the two
#: can be diffed by eye; the check in that module asserts the files exist.
SKINS = {
    "chicago": dict(font="Georgia", size=10.5, name_size=20, name_align="center",
                    heading_size=11, heading_caps=True, rule="heading", contact_align="center"),
    "zurich": dict(font="Arial", size=10, name_size=19, name_align="left",
                   heading_size=9, heading_caps=True, rule=None, heading_color="55595F"),
    "cambridge": dict(font="Garamond", size=11, name_size=21, name_align="left",
                      heading_size=11.5, heading_caps=False, rule="heading"),
    "meridian": dict(font="Calibri", size=10.5, name_size=25, name_align="left",
                     heading_size=10, heading_caps=True, rule="header", lead="company"),
    "compact": dict(font="Arial", size=9.5, name_size=16, name_align="left",
                    heading_size=8.5, heading_caps=True, rule="heading", margin=0.5),
    "ledger": dict(font="Cambria", size=10.5, name_size=18, name_align="left",
                   heading_size=9, heading_caps=True, rule=None),
    "bulletin": dict(font="Verdana", size=9.5, name_size=18, name_align="left",
                     heading_size=9, heading_caps=True, rule="heading", heading_color="4A4E55"),
}


def tag(doc, text):
    """A paragraph holding only a Jinja control tag — docxtpl removes it."""
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def bottom_border(paragraph, size=6):
    """A rule under a paragraph. python-docx has no API for it, so the border
    element is added directly; this is the standard incantation."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)


def build(template_id: str, skin: dict) -> Path:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = skin["font"]
    normal.font.size = Pt(skin["size"])
    # Word ignores the style font for East Asian runs unless told twice.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), skin["font"])

    margin = skin.get("margin", 0.6)
    for section in doc.sections:
        section.top_margin = Pt(margin * 72)
        section.bottom_margin = Pt(margin * 72)
        section.left_margin = Pt((margin + 0.1) * 72)
        section.right_margin = Pt((margin + 0.1) * 72)

    # ---- header. In the body, never in a Word header: parsers skip those.
    name_p = doc.add_paragraph()
    if skin["name_align"] == "center":
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run("{{ name }}")
    name_run.bold = skin["name_size"] < 24
    name_run.font.size = Pt(skin["name_size"])

    contact_p = doc.add_paragraph()
    if skin.get("contact_align") == "center":
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # The same middot-joined line as the HTML, and the same reason: an icon
    # reaches a parser as an unrecognised glyph attached to nothing.
    contact_run = contact_p.add_run(
        "{{ contact.email }}{% if contact.phone %} · {{ contact.phone }}{% endif %}"
        "{% if contact.location %} · {{ contact.location }}{% endif %}"
        "{% if contact.linkedin %} · {{ contact.linkedin }}{% endif %}"
    )
    contact_run.font.size = Pt(skin["size"] - 1)
    if skin.get("rule") == "header":
        bottom_border(contact_p, size=12)

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if skin["heading_caps"] else text)
        run.bold = True
        run.font.size = Pt(skin["heading_size"])
        if skin.get("heading_color"):
            run.font.color.rgb = RGBColor.from_string(skin["heading_color"])
        if skin.get("rule") == "heading":
            bottom_border(p)
        return p

    # ---- summary
    tag(doc, "{% if summary %}")
    heading("Professional Summary")
    doc.add_paragraph().add_run("{{ summary }}")
    tag(doc, "{% endif %}")

    # ---- skills
    tag(doc, "{% if skills %}")
    heading("Skills")
    doc.add_paragraph().add_run('{{ skills | join(" · ") }}')
    tag(doc, "{% endif %}")

    # ---- experience. Standard heading; a parser categorises by this string.
    tag(doc, "{% if roles %}")
    heading("Work Experience")
    tag(doc, "{% for role in roles %}")

    lead, second = (
        ("{{ role.company }}", "{{ role.title }}")
        if skin.get("lead") == "company"
        else ("{{ role.title }}", "{{ role.company }}")
    )
    head_p = doc.add_paragraph()
    head_run = head_p.add_run(lead)
    head_run.bold = True
    head_p.add_run("\t")
    head_p.add_run("{{ role.start | when }} – {{ role.end | when }}")

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(second + "{% if role.location %} — {{ role.location }}{% endif %}")
    sub_run.italic = skin["font"] in ("Georgia", "Garamond", "Cambria")

    tag(doc, "{% for bullet in role.bullets %}")
    doc.add_paragraph(style="List Bullet").add_run("{{ bullet }}")
    tag(doc, "{% endfor %}")
    tag(doc, "{% endfor %}")
    tag(doc, "{% endif %}")

    # ---- projects
    tag(doc, "{% if projects %}")
    heading("Projects")
    tag(doc, "{% for project in projects %}")
    proj_p = doc.add_paragraph()
    proj_p.add_run("{{ project.name }}").bold = True
    proj_p.add_run("{% if project.stack %} — {{ project.stack }}{% endif %}")
    tag(doc, "{% if project.description %}")
    doc.add_paragraph().add_run("{{ project.description }}")
    tag(doc, "{% endif %}")
    tag(doc, "{% endfor %}")
    tag(doc, "{% endif %}")

    # ---- education. `institution`, not `school`: see the module docstring.
    tag(doc, "{% if education %}")
    heading("Education")
    tag(doc, "{% for entry in education %}")
    edu_p = doc.add_paragraph()
    edu_p.add_run("{{ entry.degree }}").bold = True
    edu_p.add_run("\t")
    edu_p.add_run("{{ entry.year | when }}")
    tag(doc, "{% if entry.institution %}")
    doc.add_paragraph().add_run("{{ entry.institution }}")
    tag(doc, "{% endif %}")
    tag(doc, "{% endfor %}")
    tag(doc, "{% endif %}")

    # ---- certifications
    tag(doc, "{% if certifications %}")
    heading("Certifications")
    tag(doc, "{% for cert in certifications %}")
    doc.add_paragraph(style="List Bullet").add_run(
        "{{ cert.name if cert is mapping else cert }}"
        "{% if cert is mapping and cert.year %} ({{ cert.year | when }}){% endif %}"
    )
    tag(doc, "{% endfor %}")
    tag(doc, "{% endif %}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"{template_id}.docx"
    doc.save(str(path))
    return path


if __name__ == "__main__":
    for template_id, skin in SKINS.items():
        path = build(template_id, skin)
        print(f"wrote {path.relative_to(Path(__file__).resolve().parent.parent)}")
