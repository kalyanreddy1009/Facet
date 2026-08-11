"""One-time build script for templates/resume_template.docx.

Not part of the running app. The spec calls for this template to be "built
once in Word/LibreOffice with docxtpl placeholders typed directly into it."
No interactive Word/LibreOffice is available in this environment, so it's
built here via python-docx instead — each `add_run(...)` call places a
placeholder in its own single run, which sidesteps docxtpl's classic
gotcha (Word splitting a `{{ tag }}` across multiple runs when a human
types it, especially with autocorrect on). Re-run this script any time the
template needs to change; treat the resulting .docx as the fixed, sacred
shell afterward — nothing in the pipeline regenerates it at request time.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT_PATH = Path(__file__).resolve().parent / "resume_template.docx"


def add_tag_paragraph(doc, tag):
    """A paragraph containing only a Jinja control tag — docxtpl removes it."""
    p = doc.add_paragraph()
    p.add_run(tag)
    return p


def set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(10.5)


def build():
    doc = Document()
    set_base_font(doc)

    for section in doc.sections:
        section.top_margin = Pt(0.6 * 72)
        section.bottom_margin = Pt(0.6 * 72)
        section.left_margin = Pt(0.7 * 72)
        section.right_margin = Pt(0.7 * 72)

    name_p = doc.add_paragraph()
    name_run = name_p.add_run("{{ name }}")
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(
        "{{ contact.email }} · {{ contact.phone }} · {{ contact.location }} · {{ contact.linkedin }}"
    )
    contact_run.font.size = Pt(9.5)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11.5)

    # Summary
    add_tag_paragraph(doc, "{% if summary %}")
    add_heading("Summary")
    doc.add_paragraph().add_run("{{ summary }}")
    add_tag_paragraph(doc, "{% endif %}")

    # Skills
    add_tag_paragraph(doc, "{% if skills %}")
    add_heading("Skills")
    doc.add_paragraph().add_run('{{ skills | join(" · ") }}')
    add_tag_paragraph(doc, "{% endif %}")

    # Experience
    add_tag_paragraph(doc, "{% if roles %}")
    add_heading("Experience")
    add_tag_paragraph(doc, "{% for role in roles %}")

    role_header = doc.add_paragraph()
    r1 = role_header.add_run("{{ role.title }}")
    r1.bold = True
    role_header.add_run("\t")
    r2 = role_header.add_run("{{ role.start }} – {{ role.end }}")
    r2.bold = True

    role_sub = doc.add_paragraph()
    sub_run = role_sub.add_run("{{ role.company }} - {{ role.location }}")
    sub_run.italic = True

    add_tag_paragraph(doc, "{% for bullet in role.bullets %}")
    bullet_p = doc.add_paragraph(style="List Bullet")
    bullet_p.add_run("{{ bullet }}")
    add_tag_paragraph(doc, "{% endfor %}")
    add_tag_paragraph(doc, "{% endfor %}")
    add_tag_paragraph(doc, "{% endif %}")

    # Projects — fully static, never touched by agy
    add_tag_paragraph(doc, "{% if projects %}")
    add_heading("Projects")
    add_tag_paragraph(doc, "{% for project in projects %}")
    proj_p = doc.add_paragraph()
    proj_name_run = proj_p.add_run("{{ project.name }}")
    proj_name_run.bold = True
    doc.add_paragraph().add_run("{{ project.description }}")
    add_tag_paragraph(doc, "{% endfor %}")
    add_tag_paragraph(doc, "{% endif %}")

    # Education
    add_tag_paragraph(doc, "{% if education %}")
    add_heading("Education")
    add_tag_paragraph(doc, "{% for edu in education %}")
    edu_p = doc.add_paragraph()
    edu_r1 = edu_p.add_run("{{ edu.school }}")
    edu_r1.bold = True
    edu_p.add_run("\t")
    edu_r2 = edu_p.add_run("{{ edu.year }}")
    edu_r2.bold = True
    doc.add_paragraph().add_run("{{ edu.degree }}")
    add_tag_paragraph(doc, "{% endfor %}")
    add_tag_paragraph(doc, "{% endif %}")

    # Certifications
    add_tag_paragraph(doc, "{% if certifications %}")
    add_heading("Certifications")
    add_tag_paragraph(doc, "{% for cert in certifications %}")
    doc.add_paragraph(style="List Bullet").add_run("{{ cert }}")
    add_tag_paragraph(doc, "{% endfor %}")
    add_tag_paragraph(doc, "{% endif %}")

    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
